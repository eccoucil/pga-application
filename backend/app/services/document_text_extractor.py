"""
Fallback document text extraction using Python-native libraries.

No external API keys required. Handles PDF, DOCX, TXT, XLSX, CSV.
Used when LlamaExtract (LLAMA_CLOUD_API_KEY) is not configured.
"""

import csv
import logging
import os
from dataclasses import dataclass, field
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class ExtractionResult:
    """Result of text extraction from a document."""

    text: str
    page_count: int = 0
    word_count: int = 0
    format: str = "unknown"
    title: Optional[str] = None
    metadata: dict = field(default_factory=dict)


class DocumentTextExtractor:
    """
    Pure Python document text extractor.

    Supports PDF, DOCX, XLSX, TXT, CSV without external API keys.
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".xlsx", ".xls", ".csv"}

    async def extract_text(self, file_path: str, filename: str) -> ExtractionResult:
        """
        Extract text from a document file.

        Args:
            file_path: Path to the temporary file on disk
            filename: Original filename (used for extension detection)

        Returns:
            ExtractionResult with extracted text and metadata
        """
        ext = os.path.splitext(filename)[1].lower()

        if ext == ".pdf":
            return self._extract_pdf(file_path, filename)
        elif ext == ".docx":
            return self._extract_docx(file_path, filename)
        elif ext in (".xlsx", ".xls"):
            return self._extract_xlsx(file_path, filename)
        elif ext == ".csv":
            return self._extract_csv(file_path, filename)
        elif ext == ".txt":
            return self._extract_txt(file_path, filename)
        else:
            # Attempt plain text as last resort
            logger.warning(f"Unsupported extension {ext}, attempting plain text read")
            return self._extract_txt(file_path, filename)

    def _extract_pdf(self, file_path: str, filename: str) -> ExtractionResult:
        """Extract text from PDF using pypdf."""
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)

        full_text = "\n\n".join(pages)
        title = reader.metadata.title if reader.metadata else None

        return ExtractionResult(
            text=full_text,
            page_count=len(reader.pages),
            word_count=len(full_text.split()),
            format="pdf",
            title=title or os.path.splitext(filename)[0],
            metadata={
                "author": reader.metadata.author if reader.metadata else None,
                "subject": reader.metadata.subject if reader.metadata else None,
            },
        )

    def _extract_docx(self, file_path: str, filename: str) -> ExtractionResult:
        """Extract text from DOCX using python-docx."""
        from docx import Document

        doc = Document(file_path)

        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        full_text = "\n\n".join(paragraphs)

        # Extract title from core properties
        title = None
        if doc.core_properties and doc.core_properties.title:
            title = doc.core_properties.title

        return ExtractionResult(
            text=full_text,
            page_count=1,  # DOCX doesn't expose page count natively
            word_count=len(full_text.split()),
            format="docx",
            title=title or os.path.splitext(filename)[0],
            metadata={
                "author": (doc.core_properties.author if doc.core_properties else None),
            },
        )

    def _extract_xlsx(self, file_path: str, filename: str) -> ExtractionResult:
        """Extract text from XLSX using openpyxl."""
        from openpyxl import load_workbook

        wb = load_workbook(file_path, read_only=True, data_only=True)
        sheets_text = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                row_vals = [str(v) for v in row if v is not None]
                if row_vals:
                    rows.append(" | ".join(row_vals))
            if rows:
                sheets_text.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows))

        wb.close()
        full_text = "\n\n".join(sheets_text)

        return ExtractionResult(
            text=full_text,
            page_count=len(wb.sheetnames),
            word_count=len(full_text.split()),
            format="xlsx",
            title=os.path.splitext(filename)[0],
        )

    def _extract_csv(self, file_path: str, filename: str) -> ExtractionResult:
        """Extract text from CSV."""
        rows = []
        with open(file_path, newline="", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            for row in reader:
                row_text = " | ".join(cell.strip() for cell in row if cell.strip())
                if row_text:
                    rows.append(row_text)

        full_text = "\n".join(rows)

        return ExtractionResult(
            text=full_text,
            page_count=1,
            word_count=len(full_text.split()),
            format="csv",
            title=os.path.splitext(filename)[0],
        )

    def _extract_txt(self, file_path: str, filename: str) -> ExtractionResult:
        """Extract text from plain text file."""
        with open(file_path, encoding="utf-8", errors="replace") as f:
            full_text = f.read()

        return ExtractionResult(
            text=full_text,
            page_count=1,
            word_count=len(full_text.split()),
            format="txt",
            title=os.path.splitext(filename)[0],
        )
